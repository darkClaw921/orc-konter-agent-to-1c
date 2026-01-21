"""
Интеграционные тесты для API эндпоинтов
"""
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

from app.models.database import Contract, ContractData, ProcessingState
from app.models.enums import LegalEntityType


def test_contract_upload_endpoint(client, tmp_path, auth_headers):
    """
    Тест загрузки контракта через POST /api/v1/contracts/upload
    Создание тестового DOCX файла, загрузка, проверка статуса 200 и response.json()['status'] == 'processing'
    """
    from docx import Document
    
    # Создаем тестовый DOCX файл
    doc = Document()
    doc.add_heading('Тестовый договор', 0)
    doc.add_paragraph('ООО "Тест", ИНН 1234567890')
    
    test_file_path = tmp_path / "test_contract.docx"
    doc.save(str(test_file_path))
    
    # Загружаем файл через API
    with open(test_file_path, "rb") as f:
        files = {"file": ("test_contract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = client.post(
            "/api/v1/contracts/upload",
            files=files,
            headers=auth_headers
        )
    
    # Проверяем ответ
    assert response.status_code == 201, f"Expected 201, got {response.status_code}: {response.text}"
    
    data = response.json()
    assert "contract_id" in data, "Response should contain 'contract_id'"
    assert "status" in data, "Response should contain 'status'"
    assert data["status"] == "uploaded", f"Expected status 'uploaded', got '{data['status']}'"
    assert "filename" in data, "Response should contain 'filename'"
    assert data["filename"] == "test_contract.docx", f"Expected filename 'test_contract.docx', got '{data['filename']}'"


def test_get_contract_status(client, db_session, auth_headers):
    """
    Тест получения статуса контракта через GET /api/v1/contracts/{contract_id}/status
    Загрузка контракта, получение статуса, проверка наличия полей
    """
    # Создаем тестовый контракт в БД
    contract = Contract(
        original_filename="test_contract.docx",
        file_path="/tmp/test_contract.docx",
        file_size_bytes=1024,
        status=ProcessingState.PROCESSING
    )
    db_session.add(contract)
    db_session.commit()
    db_session.refresh(contract)
    
    # Получаем статус через API
    response = client.get(
        f"/api/v1/contracts/{contract.id}/status",
        headers=auth_headers
    )
    
    # Проверяем ответ
    assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
    
    data = response.json()
    assert "contract_id" in data, "Response should contain 'contract_id'"
    assert "status" in data, "Response should contain 'status'"
    assert data["contract_id"] == contract.id, f"Expected contract_id {contract.id}, got {data['contract_id']}"
    assert data["status"] == "processing", f"Expected status 'processing', got '{data['status']}'"
    assert "created_at" in data, "Response should contain 'created_at'"


def test_validation_api(client, auth_headers):
    """
    Тест валидации данных через POST /api/v1/testing/validate-extraction
    Отправка тестовых данных, проверка что is_valid == True
    """
    # Тестовые данные для валидации
    test_data = {
        "inn": "7707083893",
        "full_name": "Общество с ограниченной ответственностью 'Пример'",
        "short_name": "ООО 'Пример'",
        "organizational_form": "ООО",
        "legal_entity_type": "Юридическое лицо",
        "role": "Поставщик",
        "kpp": "770701001",
        "contract_price": "1000000.00"
    }
    
    # Отправляем запрос на валидацию
    response = client.post(
        "/api/v1/testing/validate-extraction",
        json={"extracted_data": test_data},
        headers=auth_headers
    )
    
    # Проверяем ответ
    assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
    
    data = response.json()
    assert "is_valid" in data, "Response should contain 'is_valid'"
    assert data["is_valid"] is True, f"Expected is_valid=True, got {data['is_valid']}"
    assert "errors" in data, "Response should contain 'errors'"
    assert "warnings" in data, "Response should contain 'warnings'"
    assert "validated_data" in data, "Response should contain 'validated_data'"


def test_get_contract_data(client, db_session, auth_headers):
    """
    Дополнительный тест для получения данных контракта
    """
    # Создаем тестовый контракт с данными
    contract = Contract(
        original_filename="test_contract.docx",
        file_path="/tmp/test_contract.docx",
        file_size_bytes=1024,
        status=ProcessingState.COMPLETED
    )
    db_session.add(contract)
    db_session.commit()
    db_session.refresh(contract)
    
    # Создаем данные контракта
    contract_data = ContractData(
        contract_id=contract.id,
        inn="7707083893",
        kpp="770701001",
        legal_entity_type=LegalEntityType.LEGAL_ENTITY,
        full_name="ООО 'Пример'",
        short_name="ООО 'Пример'",
        contract_price=1000000.00
    )
    db_session.add(contract_data)
    db_session.commit()
    
    # Получаем данные через API
    response = client.get(
        f"/api/v1/contracts/{contract.id}/data",
        headers=auth_headers
    )
    
    # Проверяем ответ
    assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
    
    data = response.json()
    assert "contract_id" in data, "Response should contain 'contract_id'"
    assert "inn" in data, "Response should contain 'inn'"
    assert data["inn"] == "7707083893", f"Expected INN 7707083893, got {data['inn']}"


def test_list_contracts(client, db_session, auth_headers):
    """
    Дополнительный тест для получения списка контрактов
    """
    # Создаем несколько тестовых контрактов
    for i in range(3):
        contract = Contract(
            original_filename=f"test_contract_{i}.docx",
            file_path=f"/tmp/test_contract_{i}.docx",
            file_size_bytes=1024,
            status=ProcessingState.UPLOADED
        )
        db_session.add(contract)
    
    db_session.commit()
    
    # Получаем список контрактов через API
    response = client.get(
        "/api/v1/contracts/",
        headers=auth_headers
    )
    
    # Проверяем ответ
    assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
    
    data = response.json()
    assert "contracts" in data, "Response should contain 'contracts'"
    assert "total" in data, "Response should contain 'total'"
    assert len(data["contracts"]) >= 3, f"Expected at least 3 contracts, got {len(data['contracts'])}"


def test_delete_contract(client, db_session, auth_headers, tmp_path):
    """
    Дополнительный тест для удаления контракта
    """
    # Создаем тестовый файл
    test_file = tmp_path / "test_contract.docx"
    test_file.write_bytes(b"test content")
    
    # Создаем тестовый контракт
    contract = Contract(
        original_filename="test_contract.docx",
        file_path=str(test_file),
        file_size_bytes=len(b"test content"),
        status=ProcessingState.UPLOADED
    )
    db_session.add(contract)
    db_session.commit()
    db_session.refresh(contract)
    
    # Удаляем контракт через API
    response = client.delete(
        f"/api/v1/contracts/{contract.id}",
        headers=auth_headers
    )
    
    # Проверяем ответ
    assert response.status_code == 204, f"Expected 204, got {response.status_code}: {response.text}"
    
    # Проверяем что контракт удален из БД
    deleted_contract = db_session.query(Contract).filter(Contract.id == contract.id).first()
    assert deleted_contract is None, "Contract should be deleted from database"
