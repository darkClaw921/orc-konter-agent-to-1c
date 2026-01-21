"""
Конфигурация и фикстуры для тестов
"""
import os
import tempfile
from pathlib import Path
from typing import Generator

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.main import app
from app.models.database import Base, get_db, SessionLocal
from app.core.security import create_access_token


# Тестовая БД в памяти
TEST_DATABASE_URL = "sqlite:///:memory:"

# Создаем engine для тестов
test_engine = create_engine(
    TEST_DATABASE_URL,
    connect_args={"check_same_thread": False},
    poolclass=StaticPool,
)

TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)


@pytest.fixture(scope="function")
def db_session() -> Generator:
    """Фикстура для тестовой сессии БД"""
    # Создаем все таблицы
    Base.metadata.create_all(bind=test_engine)
    
    # Создаем сессию
    session = TestingSessionLocal()
    
    try:
        yield session
    finally:
        session.close()
        # Удаляем все таблицы
        Base.metadata.drop_all(bind=test_engine)


@pytest.fixture(scope="function")
def client(db_session) -> Generator:
    """Фикстура для FastAPI TestClient"""
    def override_get_db():
        try:
            yield db_session
        finally:
            pass
    
    app.dependency_overrides[get_db] = override_get_db
    
    with TestClient(app) as test_client:
        yield test_client
    
    app.dependency_overrides.clear()


@pytest.fixture(scope="function")
def auth_headers() -> dict:
    """Фикстура для заголовков авторизации"""
    token = create_access_token(data={"sub": "test_user", "username": "test_user"})
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture(scope="function")
def temp_dir() -> Generator:
    """Фикстура для временной директории"""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture(scope="function")
def sample_docx_file(temp_dir: Path) -> Path:
    """Фикстура для создания тестового DOCX файла"""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Договор оказания услуг', 0)
    
    doc.add_paragraph('ДОГОВОР № 123/2024')
    doc.add_paragraph('г. Москва')
    doc.add_paragraph('15 января 2024 г.')
    
    doc.add_paragraph('')
    doc.add_paragraph('ООО "Пример", ИНН 7707083893, КПП 770701001, именуемое в дальнейшем "Поставщик",')
    doc.add_paragraph('с одной стороны, и')
    doc.add_paragraph('ООО "Заказчик", ИНН 1234567890, КПП 123456789, именуемое в дальнейшем "Заказчик",')
    doc.add_paragraph('с другой стороны, заключили настоящий договор о нижеследующем:')
    
    doc.add_paragraph('')
    doc.add_paragraph('1. ПРЕДМЕТ ДОГОВОРА')
    doc.add_paragraph('1.1. Поставщик обязуется оказать Заказчику консультационные услуги.')
    
    doc.add_paragraph('')
    doc.add_paragraph('2. СТОИМОСТЬ ДОГОВОРА')
    doc.add_paragraph('2.1. Стоимость услуг составляет 1 000 000,00 (один миллион) рублей, включая НДС 20%.')
    
    doc.add_paragraph('')
    doc.add_paragraph('3. СРОКИ ОКАЗАНИЯ УСЛУГ')
    doc.add_paragraph('3.1. Услуги оказываются с 01.02.2024 по 31.12.2024.')
    
    doc.add_paragraph('')
    doc.add_paragraph('4. МЕСТО ОКАЗАНИЯ УСЛУГ')
    doc.add_paragraph('4.1. Услуги оказываются по адресу: г. Москва, ул. Примерная, д. 1')
    
    doc.add_paragraph('')
    doc.add_paragraph('5. ОТВЕТСТВЕННЫЕ ЛИЦА')
    doc.add_paragraph('5.1. От Поставщика: Иванов Иван Иванович, тел. +7 (495) 123-45-67, email: ivanov@example.com')
    doc.add_paragraph('5.2. От Заказчика: Петров Петр Петрович, тел. +7 (495) 765-43-21, email: petrov@example.com')
    
    file_path = temp_dir / "test_contract.docx"
    doc.save(str(file_path))
    
    return file_path


@pytest.fixture(scope="function")
def mock_llm_service(monkeypatch):
    """Фикстура для мокирования LLM сервиса"""
    mock_data = {
        "inn": "7707083893",
        "full_name": "Общество с ограниченной ответственностью 'Пример'",
        "short_name": "ООО 'Пример'",
        "organizational_form": "ООО",
        "legal_entity_type": "Юридическое лицо",
        "role": "Поставщик",
        "kpp": "770701001",
        "contract_name": "Договор оказания услуг",
        "contract_number": "123/2024",
        "contract_date": "2024-01-15",
        "contract_price": "1000000.00",
        "vat_percent": "20.00",
        "vat_type": "Добавляется",
        "service_description": "Оказание консультационных услуг",
        "service_start_date": "2024-02-01",
        "service_end_date": "2024-12-31",
        "locations": [
            {
                "address": "г. Москва, ул. Примерная, д. 1",
                "responsible_person": {
                    "name": "Иванов Иван Иванович",
                    "phone": "+7 (495) 123-45-67",
                    "email": "ivanov@example.com"
                }
            }
        ],
        "responsible_persons": [
            {
                "name": "Петров Петр Петрович",
                "phone": "+7 (495) 765-43-21",
                "email": "petrov@example.com"
            }
        ]
    }
    
    async def mock_extract_contract_data(self, document_text: str, retry_count: int = 3):
        return mock_data
    
    from app.services.llm_service import LLMService
    monkeypatch.setattr(LLMService, "extract_contract_data", mock_extract_contract_data)


@pytest.fixture(scope="function")
def mock_onec_service(monkeypatch):
    """Фикстура для мокирования OneC сервиса"""
    async def mock_find_counterparty_by_inn(self, inn: str):
        # По умолчанию контрагент не найден
        return None
    
    async def mock_create_counterparty(self, contract_data: dict, document_path: str):
        # Возвращаем тестовый UUID
        return "test-counterparty-uuid-12345"
    
    from app.services.oneс_service import OneCService
    monkeypatch.setattr(OneCService, "find_counterparty_by_inn", mock_find_counterparty_by_inn)
    monkeypatch.setattr(OneCService, "create_counterparty", mock_create_counterparty)


@pytest.fixture(scope="function")
def mock_redis(monkeypatch):
    """Фикстура для мокирования Redis"""
    class MockRedis:
        def __init__(self):
            self._data = {}
        
        async def get(self, key: str):
            return self._data.get(key)
        
        async def set(self, key: str, value: str, ex: int = None):
            self._data[key] = value.encode() if isinstance(value, str) else value
            return True
    
    mock_redis_instance = MockRedis()
    
    # Мокируем redis клиент в state_manager
    from app.agent import state_manager
    monkeypatch.setattr(state_manager, "redis", None)  # Отключаем Redis для тестов
    
    return mock_redis_instance
