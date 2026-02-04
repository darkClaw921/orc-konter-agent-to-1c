"""
Парсер DOCX и PDF файлов с использованием docling для извлечения таблиц
"""
import re
import unicodedata
from dataclasses import dataclass
from enum import Enum
from typing import Dict, List, Optional, Any
from pathlib import Path

import pandas as pd
from docx import Document
from docx.document import Document as DocType
from docling.document_converter import DocumentConverter
from docling_core.types.doc import TextItem, TableItem

from app.config import settings
from app.utils.logging import get_logger

logger = get_logger(__name__)


def clean_text(text: str) -> str:
    """
    Очистка текста от проблемных символов

    Args:
        text: Исходный текст

    Returns:
        Очищенный текст
    """
    if not text:
        return text

    # Нормализация Unicode (приводим к единообразной форме NFC)
    text = unicodedata.normalize('NFC', text)

    # Удаляем невидимые символы и replacement characters
    invisible_chars = [
        '\u200b',  # Zero-width space
        '\u200c',  # Zero-width non-joiner
        '\u200d',  # Zero-width joiner
        '\ufeff',  # Zero-width no-break space (BOM)
        '\u00ad',  # Soft hyphen
        '\u2060',  # Word joiner
        '\ufffd',  # Replacement character (�)
    ]
    for char in invisible_chars:
        text = text.replace(char, '')

    # Удаляем управляющие символы (кроме \n, \r, \t)
    # Категория Cc - это control characters
    text = ''.join(
        char for char in text
        if not unicodedata.category(char).startswith('C') or char in '\n\r\t'
    )

    # Заменяем множественные пробелы на одинарные
    text = re.sub(r'[ \t]+', ' ', text)

    # Заменяем множественные переносы строк на максимум два
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Удаляем пробелы в конце строк
    text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)

    return text


class ElementType(Enum):
    TEXT = "text"
    TABLE = "table"


@dataclass
class DocumentElement:
    """Элемент документа (текст или таблица)"""
    type: ElementType
    content: str
    index: int
    size: int
    is_splittable: bool = True
    table_index: Optional[int] = None
    row_count: Optional[int] = None
    column_count: Optional[int] = None


class DocumentProcessor:
    """
    Процессор для извлечения текста из DOCX и PDF документов с использованием docling
    """
    
    def __init__(self):
        self.document: Optional[DocType] = None
        self.docling_result: Optional[Any] = None
        self.docling_converter: Optional[DocumentConverter] = None
        self.file_path: Optional[str] = None
        self.file_type: Optional[str] = None  # 'docx' или 'pdf'
        self.raw_text: str = ""
        self.paragraphs: List[str] = []
        self.tables: List[Dict[str, Any]] = []  # Список словарей с данными таблиц
        self.document_elements: List[DocumentElement] = []  # Элементы документа в порядке появления
    
    def load_document(self, file_path: str) -> bool:
        """
        Загрузить DOCX или PDF документ с использованием docling
        
        Args:
            file_path: Путь к файлу документа
            
        Returns:
            True если документ успешно загружен, False в противном случае
        """
        self.file_path = file_path
        file_ext = Path(file_path).suffix.lower()
        
        # Определяем тип файла
        if file_ext == '.docx':
            self.file_type = 'docx'
        elif file_ext == '.pdf':
            self.file_type = 'pdf'
        else:
            logger.error("Unsupported file type", file_path=file_path, extension=file_ext)
            return False
        
        # Пытаемся загрузить через docling
        try:
            if not self.docling_converter:
                self.docling_converter = DocumentConverter()
            
            self.docling_result = self.docling_converter.convert(file_path)
            logger.info("Document loaded with docling", 
                       file_path=file_path, 
                       file_type=self.file_type,
                       page_count=getattr(self.docling_result.input, 'page_count', None))
            
            # Для обратной совместимости также загружаем через python-docx для DOCX
            if self.file_type == 'docx':
                try:
                    self.document = Document(file_path)
                except Exception as e:
                    logger.warning("Failed to load with python-docx, using docling only",
                                 error=str(e))
            
            return True
        except Exception as e:
            logger.error("Failed to load document with docling", 
                        error=str(e), 
                        file_path=file_path,
                        error_type=type(e).__name__)
            
            # Fallback на python-docx для DOCX файлов
            if self.file_type == 'docx':
                try:
                    logger.info("Trying fallback to python-docx")
                    self.document = Document(file_path)
                    logger.info("Document loaded with python-docx fallback", file_path=file_path)
                    return True
                except Exception as fallback_error:
                    logger.error("Fallback to python-docx also failed",
                               error=str(fallback_error))
            
            return False
    
    def extract_text(self) -> str:
        """
        Извлечь весь текст из документа включая таблицы в markdown формате

        Returns:
            Текст документа с таблицами в markdown формате
        """
        if not self.docling_result and not self.document:
            raise ValueError("Document not loaded")

        text_parts = []
        self.paragraphs = []
        self.tables = []
        self.document_elements = []
        element_index = 0

        # Используем docling если доступен
        if self.docling_result:
            doc = self.docling_result.document

            # Извлекаем текст и таблицы через docling API
            for item, _level in doc.iterate_items():
                if isinstance(item, TextItem):
                    text = clean_text(item.text.strip())
                    if text:
                        text_parts.append(text)
                        self.paragraphs.append(text)
                        # Добавляем текстовый элемент
                        self.document_elements.append(DocumentElement(
                            type=ElementType.TEXT,
                            content=text,
                            index=element_index,
                            size=len(text),
                            is_splittable=True
                        ))
                        element_index += 1
                elif isinstance(item, TableItem):
                    # Извлекаем таблицу в структурированном виде
                    table_data = self._extract_table_from_docling(item)
                    if table_data:
                        self.tables.append(table_data)
                        # Формируем markdown для таблицы
                        table_markdown = f"\n\n## Таблица {len(self.tables)}\n\n{table_data['markdown']}\n"
                        # Добавляем таблицу в markdown формате в текст
                        text_parts.append(f"\n\n## Таблица {len(self.tables)}\n\n")
                        text_parts.append(table_data['markdown'])
                        text_parts.append("\n")
                        # Добавляем табличный элемент (неразделяемый!)
                        self.document_elements.append(DocumentElement(
                            type=ElementType.TABLE,
                            content=table_markdown,
                            index=element_index,
                            size=len(table_markdown),
                            is_splittable=False,  # Таблицу нельзя разбивать!
                            table_index=table_data['index'],
                            row_count=table_data['row_count'],
                            column_count=table_data['column_count']
                        ))
                        element_index += 1

        # Fallback на python-docx для DOCX если docling не сработал
        elif self.document and self.file_type == 'docx':
            logger.info("Using python-docx fallback for text extraction")

            # Извлечение текста из параграфов
            for paragraph in self.document.paragraphs:
                if paragraph.text.strip():
                    text = clean_text(paragraph.text.strip())
                    if text:
                        text_parts.append(text)
                        self.paragraphs.append(text)
                        # Добавляем текстовый элемент
                        self.document_elements.append(DocumentElement(
                            type=ElementType.TEXT,
                            content=text,
                            index=element_index,
                            size=len(text),
                            is_splittable=True
                        ))
                        element_index += 1

            # Извлечение текста из таблиц (простой формат)
            for table_idx, table in enumerate(self.document.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(clean_text(cell.text.strip()))
                    table_data.append(row_data)

                # Сохраняем таблицу в старом формате для обратной совместимости
                table_info = {
                    "index": table_idx,
                    "markdown": self._convert_table_to_markdown(table_data),
                    "dataframe": pd.DataFrame(table_data[1:], columns=table_data[0]) if table_data else pd.DataFrame(),
                    "row_count": len(table_data),
                    "column_count": len(table_data[0]) if table_data else 0
                }
                self.tables.append(table_info)

                table_markdown = f"\n\n## Таблица {table_idx + 1}\n\n{table_info['markdown']}\n"
                text_parts.append(f"\n\n## Таблица {table_idx + 1}\n\n")
                text_parts.append(self.tables[-1]['markdown'])
                text_parts.append("\n")

                # Добавляем табличный элемент (неразделяемый!)
                self.document_elements.append(DocumentElement(
                    type=ElementType.TABLE,
                    content=table_markdown,
                    index=element_index,
                    size=len(table_markdown),
                    is_splittable=False,  # Таблицу нельзя разбивать!
                    table_index=table_info['index'],
                    row_count=table_info['row_count'],
                    column_count=table_info['column_count']
                ))
                element_index += 1

        self.raw_text = "\n".join(text_parts)
        logger.info("Text extracted",
                   length=len(self.raw_text),
                   paragraphs_count=len(self.paragraphs),
                   tables_count=len(self.tables),
                   elements_count=len(self.document_elements))
        return self.raw_text
    
    def _extract_table_from_docling(self, table_item: TableItem) -> Optional[Dict[str, Any]]:
        """
        Извлечь таблицу из docling TableItem в структурированном виде
        
        Args:
            table_item: TableItem из docling
            
        Returns:
            Словарь с данными таблицы или None если не удалось извлечь
        """
        try:
            df = table_item.export_to_dataframe(doc=self.docling_result.document)
            
            # Конвертируем DataFrame в markdown
            markdown_table = df.to_markdown(index=False)
            
            table_data = {
                "index": len(self.tables),
                "markdown": markdown_table,
                "dataframe": df,
                "row_count": len(df),
                "column_count": len(df.columns)
            }
            
            logger.debug("Table extracted", 
                        table_index=table_data["index"],
                        rows=table_data["row_count"],
                        columns=table_data["column_count"])
            
            return table_data
        except Exception as e:
            logger.warning("Failed to extract table from docling",
                         error=str(e),
                         error_type=type(e).__name__)
            return None
    
    def _convert_table_to_markdown(self, table_data: List[List[str]]) -> str:
        """
        Конвертировать таблицу в markdown формат
        
        Args:
            table_data: Список строк таблицы
            
        Returns:
            Таблица в markdown формате
        """
        if not table_data:
            return ""
        
        # Первая строка - заголовки
        headers = table_data[0] if table_data else []
        if not headers:
            return ""
        
        # Формируем markdown таблицу
        markdown_lines = []
        
        # Заголовки
        markdown_lines.append("| " + " | ".join(str(h) for h in headers) + " |")
        
        # Разделитель
        markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # Данные
        for row in table_data[1:]:
            # Дополняем строку до нужной длины если нужно
            row_padded = row + [""] * (len(headers) - len(row))
            markdown_lines.append("| " + " | ".join(str(cell) for cell in row_padded[:len(headers)]) + " |")
        
        return "\n".join(markdown_lines)
    
    def extract_tables(self) -> List[Dict[str, Any]]:
        """
        Извлечь все таблицы из документа в структурированном виде
        
        Returns:
            Список словарей с данными таблиц (markdown, DataFrame, метаданные)
        """
        if not self.tables:
            # Если таблицы еще не извлечены, вызываем extract_text
            self.extract_text()

        logger.info("Tables extracted",
                   tables_count=len(self.tables),
                   tables=self.tables)
        return self.tables
    
    def get_tables_markdown(self) -> str:
        """
        Получить все таблицы в markdown формате для включения в контекст LLM
        
        Returns:
            Все таблицы документа в markdown формате, разделенные заголовками
        """
        if not self.tables:
            self.extract_tables()
        
        if not self.tables:
            return ""
        
        markdown_parts = []
        for table_data in self.tables:
            markdown_parts.append(f"## Таблица {table_data['index'] + 1}")
            markdown_parts.append("")
            markdown_parts.append(table_data['markdown'])
            markdown_parts.append("")
        
        return "\n".join(markdown_parts)
    
    def extract_sections(self) -> Dict[str, str]:
        """Извлечь основные разделы документа"""
        if not self.raw_text:
            self.extract_text()
        
        sections = {}
        
        # Ищем типичные разделы контрактов
        section_patterns = {
            "header": r"(?:договор|контракт|соглашение)(.*?)(?:стороны|участники|предмет)",
            "parties": r"(?:стороны|участники):(.*?)(?:предмет|место оказания)",
            "subject": r"(?:предмет|услуг)(.*?)(?:стоимость|цена|размер)",
            "price": r"(?:стоимость|цена|размер вознаграждения)(.*?)(?:срок|дата|условия)",
            "terms": r"(?:срок|сроки|период)(.*?)(?:условия|порядок|оплата)",
            "conditions": r"(?:условия|прочие положения)(.*?)(?:подпись|подписано|дата)",
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, self.raw_text, re.IGNORECASE | re.DOTALL)
            if match:
                sections[section_name] = match.group(1)[:1000]
        
        logger.info("Sections extracted", count=len(sections))
        return sections
    
    def get_context_for_llm(self, max_tokens: int = 8000) -> str:
        """
        Подготовить контекст для отправки в LLM.
        Включает текст документа и таблицы в markdown формате.
        Ограничивает размер, чтобы не превысить лимиты токенов.
        
        Args:
            max_tokens: Максимальное количество токенов в контексте
            
        Returns:
            Контекст документа с таблицами в markdown формате
        """
        if not self.raw_text:
            self.extract_text()
        
        # Таблицы уже включены в raw_text через extract_text()
        # Но если нужно, можно добавить дополнительную обработку
        
        # Если текст слишком большой, берем только начало
        # Примерно 4 символа на токен
        max_chars = max_tokens * 4
        if len(self.raw_text) > max_chars:
            truncated_text = self.raw_text[:max_chars]
            logger.warning("Document truncated", 
                         original=len(self.raw_text), 
                         truncated=len(truncated_text),
                         max_tokens=max_tokens)
            return truncated_text
        
        return self.raw_text
    
    def estimate_tokens(self, text: str) -> int:
        """Оценить количество токенов в тексте"""
        # Примерная оценка: 1 токен ≈ 4 символа
        return len(text) // 4
    
    def _elements_to_text(self, elements: List[DocumentElement]) -> str:
        """
        Преобразовать список элементов в текст

        Args:
            elements: Список элементов документа

        Returns:
            Объединенный текст элементов
        """
        return "\n".join(elem.content for elem in elements)

    def _get_overlap_elements(self, elements: List[DocumentElement], max_overlap_size: int) -> List[DocumentElement]:
        """
        Получить элементы для overlap (только текстовые)

        Args:
            elements: Список элементов текущего чанка
            max_overlap_size: Максимальный размер overlap в символах

        Returns:
            Список элементов для overlap
        """
        overlap_elements = []
        overlap_size = 0

        # Идем с конца и берем только текстовые элементы
        for elem in reversed(elements):
            if elem.type == ElementType.TEXT:
                if overlap_size + elem.size <= max_overlap_size:
                    overlap_elements.insert(0, elem)
                    overlap_size += elem.size
                else:
                    break

        return overlap_elements

    def _split_large_text(self, text: str, max_size: int, overlap: int) -> List[str]:
        """
        Разбить большой текст на части по словам

        Args:
            text: Текст для разбиения
            max_size: Максимальный размер части
            overlap: Размер перекрытия

        Returns:
            Список частей текста
        """
        words = text.split()
        parts = []
        temp_chunk = []
        temp_size = 0

        for word in words:
            word_size = len(word) + 1  # +1 для пробела
            if temp_size + word_size > max_size and temp_chunk:
                parts.append(" ".join(temp_chunk))
                # Добавляем перекрытие из предыдущего чанка
                overlap_words = temp_chunk[-(overlap // 10):] if len(temp_chunk) > overlap // 10 else temp_chunk
                temp_chunk = overlap_words + [word]
                temp_size = sum(len(w) + 1 for w in temp_chunk)
            else:
                temp_chunk.append(word)
                temp_size += word_size

        if temp_chunk:
            parts.append(" ".join(temp_chunk))

        return parts

    def _extract_table_header_and_rows(self, table_markdown: str) -> tuple[str, str, List[str]]:
        """
        Извлечь заголовок, разделитель и строки из markdown таблицы

        Args:
            table_markdown: Таблица в markdown формате

        Returns:
            Кортеж (заголовок таблицы с ## Таблица X, строка заголовков с разделителем, список строк данных)
        """
        lines = table_markdown.strip().split('\n')

        # Ищем заголовок таблицы (## Таблица X)
        table_title = ""
        header_start_idx = 0

        for idx, line in enumerate(lines):
            if line.strip().startswith('## Таблица'):
                table_title = line.strip()
                header_start_idx = idx + 1
                break
            elif line.strip().startswith('|'):
                # Если нет заголовка, начинаем с первой строки таблицы
                header_start_idx = idx
                break

        # Пропускаем пустые строки после заголовка
        while header_start_idx < len(lines) and not lines[header_start_idx].strip():
            header_start_idx += 1

        if header_start_idx >= len(lines):
            return table_title, "", []

        # Строка заголовков столбцов
        header_line = lines[header_start_idx] if header_start_idx < len(lines) else ""

        # Строка-разделитель (| --- | --- |)
        separator_line = lines[header_start_idx + 1] if header_start_idx + 1 < len(lines) else ""

        # Объединяем заголовок и разделитель
        table_header = header_line + '\n' + separator_line if separator_line else header_line

        # Строки данных (начиная с индекса header_start_idx + 2)
        data_rows = lines[header_start_idx + 2:] if header_start_idx + 2 < len(lines) else []

        # Фильтруем пустые строки
        data_rows = [row for row in data_rows if row.strip()]

        return table_title, table_header, data_rows

    def _split_large_table(self, element: 'DocumentElement', max_size: int) -> List['DocumentElement']:
        """
        Разбить большую таблицу на несколько чанков с сохранением заголовков

        Args:
            element: Элемент таблицы для разбиения
            max_size: Максимальный размер чанка в символах

        Returns:
            Список элементов (частей таблицы)
        """
        table_title, table_header, data_rows = self._extract_table_header_and_rows(element.content)

        if not data_rows:
            # Если нет строк данных, возвращаем как есть
            return [element]

        # Размер заголовка (заголовок таблицы + заголовок столбцов + разделитель)
        header_with_title = f"{table_title}\n\n{table_header}\n" if table_title else f"{table_header}\n"
        header_size = len(header_with_title)

        # Проверяем, что заголовок сам по себе не больше лимита
        if header_size >= max_size * 0.8:
            logger.warning("Table header is too large, returning as single chunk",
                         table_index=element.table_index,
                         header_size=header_size,
                         max_size=max_size)
            return [element]

        # Доступный размер для строк данных
        available_size = max_size - header_size - 100  # Оставляем запас

        chunks = []
        current_rows = []
        current_size = 0
        part_num = 1

        for row in data_rows:
            row_size = len(row) + 1  # +1 для переноса строки

            if current_size + row_size > available_size and current_rows:
                # Создаем чанк из текущих строк
                chunk_content = header_with_title + '\n'.join(current_rows) + '\n'

                # Добавляем информацию о части
                if table_title:
                    chunk_content = chunk_content.replace(
                        table_title,
                        f"{table_title} (часть {part_num})"
                    )

                chunks.append(DocumentElement(
                    type=ElementType.TABLE,
                    content=chunk_content,
                    index=element.index,
                    size=len(chunk_content),
                    is_splittable=False,  # Уже разбито, дальше не разбивать
                    table_index=element.table_index,
                    row_count=len(current_rows),
                    column_count=element.column_count
                ))

                current_rows = [row]
                current_size = row_size
                part_num += 1
            else:
                current_rows.append(row)
                current_size += row_size

        # Добавляем последний чанк
        if current_rows:
            chunk_content = header_with_title + '\n'.join(current_rows) + '\n'

            if table_title and part_num > 1:
                chunk_content = chunk_content.replace(
                    table_title,
                    f"{table_title} (часть {part_num})"
                )

            chunks.append(DocumentElement(
                type=ElementType.TABLE,
                content=chunk_content,
                index=element.index,
                size=len(chunk_content),
                is_splittable=False,
                table_index=element.table_index,
                row_count=len(current_rows),
                column_count=element.column_count
            ))

        logger.info("Large table split into parts",
                   table_index=element.table_index,
                   original_size=element.size,
                   parts_count=len(chunks),
                   total_rows=len(data_rows))

        return chunks

    def _build_elements_from_paragraphs(self) -> List[DocumentElement]:
        """
        Fallback: построить элементы из параграфов (для обратной совместимости)

        Returns:
            Список элементов документа
        """
        elements = []
        for idx, para in enumerate(self.paragraphs):
            if para.strip():
                elements.append(DocumentElement(
                    type=ElementType.TEXT,
                    content=para.strip(),
                    index=idx,
                    size=len(para.strip()),
                    is_splittable=True
                ))
        return elements

    def split_into_chunks(self, max_chunk_size: int, overlap: int = 200, max_table_chunk_size: Optional[int] = None) -> List[str]:
        """
        Разбить текст документа на чанки с перекрытием.
        Использует document_elements для сохранения таблиц целиком.

        Args:
            max_chunk_size: Максимальный размер чанка в символах
            overlap: Размер перекрытия между чанками в символах
            max_table_chunk_size: Максимальный размер чанка для таблиц в символах (если None, используется max_chunk_size)

        Returns:
            Список текстовых чанков
        """
        if not self.raw_text:
            self.extract_text()

        if len(self.raw_text) <= max_chunk_size:
            return [self.raw_text]

        # Используем document_elements или fallback на paragraphs
        elements = self.document_elements if self.document_elements else self._build_elements_from_paragraphs()

        # Если элементов нет, разбиваем raw_text по предложениям
        if not elements:
            sentences = re.split(r'(?<=[.!?])\s+', self.raw_text)
            elements = [
                DocumentElement(
                    type=ElementType.TEXT,
                    content=s.strip(),
                    index=idx,
                    size=len(s.strip()),
                    is_splittable=True
                )
                for idx, s in enumerate(sentences) if s.strip()
            ]

        # Используем отдельный лимит для таблиц если указан
        table_max_size = max_table_chunk_size if max_table_chunk_size is not None else max_chunk_size

        chunks = []
        current_elements: List[DocumentElement] = []
        current_size = 0

        for element in elements:
            # Для неразделяемых элементов (таблиц) особая логика
            if not element.is_splittable:
                # Если таблица больше table_max_size, разбиваем с сохранением заголовков
                if element.size > table_max_size:
                    # Сохраняем текущий чанк если есть
                    if current_elements:
                        chunks.append(self._elements_to_text(current_elements))
                        current_elements = []
                        current_size = 0

                    # Разбиваем большую таблицу на части с заголовками
                    logger.info("Splitting large table into parts with headers",
                               table_index=element.table_index,
                               table_size=element.size,
                               max_table_chunk_size=table_max_size)

                    table_parts = self._split_large_table(element, table_max_size)

                    # Добавляем каждую часть как отдельный чанк
                    for part in table_parts:
                        if part.size > table_max_size:
                            # Если часть все еще слишком большая, логируем предупреждение
                            logger.warning("Table part still exceeds max_table_chunk_size",
                                         table_index=element.table_index,
                                         part_size=part.size,
                                         max_table_chunk_size=table_max_size)
                        chunks.append(part.content)
                    continue

                # Если таблица не поместится в текущий чанк, начинаем новый
                # Для таблиц используем table_max_size
                if current_size + element.size + 1 > table_max_size:
                    if current_elements:
                        chunks.append(self._elements_to_text(current_elements))
                        # Получаем overlap только из текстовых элементов
                        overlap_elements = self._get_overlap_elements(current_elements, overlap)
                        current_elements = overlap_elements.copy()
                        current_size = sum(e.size for e in current_elements)

                # Добавляем таблицу целиком
                current_elements.append(element)
                current_size += element.size + 1
                continue

            # Для разделяемых элементов (текст)
            elem_content = element.content.strip()
            if not elem_content:
                continue

            elem_size = len(elem_content)

            # Если элемент сам по себе больше max_chunk_size, разбиваем его
            if elem_size > max_chunk_size:
                # Сохраняем текущий чанк если есть
                if current_elements:
                    chunks.append(self._elements_to_text(current_elements))
                    current_elements = []
                    current_size = 0

                # Разбиваем большой текст на части
                text_parts = self._split_large_text(elem_content, max_chunk_size, overlap)
                for part in text_parts[:-1]:
                    chunks.append(part)

                # Последнюю часть оставляем как начало следующего чанка
                if text_parts:
                    current_elements = [DocumentElement(
                        type=ElementType.TEXT,
                        content=text_parts[-1],
                        index=element.index,
                        size=len(text_parts[-1]),
                        is_splittable=True
                    )]
                    current_size = len(text_parts[-1])
                continue

            # Проверяем, поместится ли элемент в текущий чанк
            if current_size + elem_size + 1 <= max_chunk_size:
                current_elements.append(DocumentElement(
                    type=ElementType.TEXT,
                    content=elem_content,
                    index=element.index,
                    size=elem_size,
                    is_splittable=True
                ))
                current_size += elem_size + 1
            else:
                # Сохраняем текущий чанк
                if current_elements:
                    chunks.append(self._elements_to_text(current_elements))

                # Начинаем новый чанк с перекрытием
                if overlap > 0 and current_elements:
                    overlap_elements = self._get_overlap_elements(current_elements, overlap)
                    overlap_size = sum(e.size for e in overlap_elements)
                    if overlap_size < max_chunk_size // 2:
                        current_elements = overlap_elements + [DocumentElement(
                            type=ElementType.TEXT,
                            content=elem_content,
                            index=element.index,
                            size=elem_size,
                            is_splittable=True
                        )]
                    else:
                        current_elements = [DocumentElement(
                            type=ElementType.TEXT,
                            content=elem_content,
                            index=element.index,
                            size=elem_size,
                            is_splittable=True
                        )]
                    current_size = sum(e.size for e in current_elements)
                else:
                    current_elements = [DocumentElement(
                        type=ElementType.TEXT,
                        content=elem_content,
                        index=element.index,
                        size=elem_size,
                        is_splittable=True
                    )]
                    current_size = elem_size

        # Добавляем последний чанк
        if current_elements:
            chunks.append(self._elements_to_text(current_elements))

        logger.info("Document split into chunks",
                   total_chunks=len(chunks),
                   avg_chunk_size=sum(len(c) for c in chunks) // len(chunks) if chunks else 0,
                   tables_count=len(self.tables),
                   elements_count=len(elements))

        return chunks
    
    def get_chunks_for_llm(self, max_tokens_per_chunk: Optional[int] = None) -> List[str]:
        """
        Получить список чанков документа для обработки через LLM

        Args:
            max_tokens_per_chunk: Максимальное количество токенов в одном чанке.
                                  Если не указано, берется из settings.MAX_CHUNK_TOKENS

        Returns:
            Список текстовых чанков
        """
        # Используем настройки из конфига если не указано явно
        if max_tokens_per_chunk is None:
            max_tokens_per_chunk = settings.MAX_CHUNK_TOKENS

        # Примерно 4 символа на токен
        max_chunk_size = max_tokens_per_chunk * 4

        # Отдельный лимит для таблиц
        max_table_chunk_size = settings.MAX_TABLE_CHUNK_TOKENS * 4

        # Перекрытие из настроек (в символах)
        overlap = settings.CHUNK_OVERLAP_TOKENS * 4

        logger.info("Preparing chunks for LLM",
                   max_tokens=max_tokens_per_chunk,
                   max_chunk_size_chars=max_chunk_size,
                   max_table_chunk_size_chars=max_table_chunk_size,
                   overlap_chars=overlap)

        return self.split_into_chunks(max_chunk_size=max_chunk_size, overlap=overlap, max_table_chunk_size=max_table_chunk_size)